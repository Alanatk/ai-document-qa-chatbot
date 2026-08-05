"""
Utility functions for file handling, multi-format document loading,
metadata preservation, text splitting, statistical calculations,
and audio speech-to-text transcription.
"""

import os
import tempfile
from typing import Any, Dict, List
import docx
from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter


def save_uploaded_file_temp(uploaded_file: Any) -> str:
    """Saves a Streamlit UploadedFile object to a temporary file on disk and

    returns its path.
    """
    _, extension = os.path.splitext(uploaded_file.name)
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=extension)
    temp_file.write(uploaded_file.read())
    temp_path = temp_file.name
    temp_file.close()
    return temp_path


def load_document_by_type(
    file_path: str, original_filename: str
) -> List[Document]:
    """Loads text content from PDF, DOCX, or TXT files into LangChain Document

    objects, ensuring page numbers and document sources are preserved in
    metadata.
    """
    _, extension = os.path.splitext(original_filename)
    extension = extension.lower()
    documents: List[Document] = []

    if extension == ".pdf":
        loader = PyPDFLoader(file_path)
        raw_docs = loader.load()
        for idx, doc in enumerate(raw_docs):
            page_num = doc.metadata.get("page", idx)
            if isinstance(page_num, int):
                page_num += 1  # 0-indexed to 1-indexed
            doc.metadata["source"] = original_filename
            doc.metadata["page"] = page_num
            documents.append(doc)

    elif extension == ".docx":
        doc_obj = docx.Document(file_path)
        paragraphs = [p.text for p in doc_obj.paragraphs if p.text.strip()]
        full_text = "\n\n".join(paragraphs)
        documents.append(
            Document(
                page_content=full_text,
                metadata={"source": original_filename, "page": 1},
            )
        )

    elif extension == ".txt":
        try:
            loader = TextLoader(file_path, encoding="utf-8")
            raw_docs = loader.load()
        except UnicodeDecodeError:
            loader = TextLoader(file_path, encoding="latin-1")
            raw_docs = loader.load()

        for doc in raw_docs:
            doc.metadata["source"] = original_filename
            doc.metadata["page"] = 1
            documents.append(doc)

    else:
        raise ValueError(f"Unsupported file format: {extension}")

    return documents


def split_documents_into_chunks(
    documents: List[Document], chunk_size: int = 1000, chunk_overlap: int = 200
) -> List[Document]:
    """Splits loaded Document objects into manageable chunks using

    RecursiveCharacterTextSplitter while maintaining original metadata.
    """
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["\n\n", "\n", ". ", " ", ""],
    )
    chunks = text_splitter.split_documents(documents)
    return chunks


def process_uploaded_file(
    uploaded_file: Any, chunk_size: int = 1000, chunk_overlap: int = 200
) -> List[Document]:
    """Full lifecycle handler for processing an uploaded Streamlit file:

    1. Saves file temporarily.
    2. Loads documents with metadata preservation.
    3. Splits text into chunks.
    4. Safely removes temporary file.
    """
    temp_path = save_uploaded_file_temp(uploaded_file)
    try:
        documents = load_document_by_type(temp_path, uploaded_file.name)
        chunks = split_documents_into_chunks(
            documents, chunk_size=chunk_size, chunk_overlap=chunk_overlap
        )
        return chunks
    finally:
        if os.path.exists(temp_path):
            try:
                os.remove(temp_path)
            except Exception:
                pass


def calculate_document_stats(
    chunks: List[Document], uploaded_files: List[Any] = None
) -> Dict[str, Any]:
    """Calculates comprehensive document processing statistics including total

    file size for UI metrics.
    """
    if not chunks:
        return {
            "total_documents": 0,
            "total_chunks": 0,
            "total_words": 0,
            "total_characters": 0,
            "total_size_mb": 0.0,
            "file_names": [],
        }

    sources = set()
    total_chars = 0
    total_words = 0

    for chunk in chunks:
        sources.add(chunk.metadata.get("source", "Unknown"))
        text = chunk.page_content
        total_chars += len(text)
        total_words += len(text.split())

    total_bytes = 0
    if uploaded_files:
        for f in uploaded_files:
            if hasattr(f, "size"):
                total_bytes += f.size

    total_size_mb = round(total_bytes / (1024 * 1024), 2)

    return {
        "total_documents": len(sources),
        "total_chunks": len(chunks),
        "total_words": total_words,
        "total_characters": total_chars,
        "total_size_mb": total_size_mb,
        "file_names": sorted(list(sources)),
    }


def transcribe_audio_bytes(audio_bytes: bytes) -> str:
    """Transcribes audio bytes into text using SpeechRecognition."""
    try:
        import speech_recognition as sr
    except ImportError:
        return ""

    temp_audio = tempfile.NamedTemporaryFile(delete=False, suffix=".wav")
    temp_audio.write(audio_bytes)
    temp_path = temp_audio.name
    temp_audio.close()

    try:
        recognizer = sr.Recognizer()
        with sr.AudioFile(temp_path) as source:
            audio_data = recognizer.record(source)
            text = recognizer.recognize_google(audio_data)
            return text
    except Exception:
        return ""
    finally:
        if os.path.exists(temp_path):
            try:
                os.remove(temp_path)
            except Exception:
                pass
